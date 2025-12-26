import io
import re
from datetime import datetime
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, white
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Flowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter(prefix="/api/xray", tags=["downloads"])

class DownloadRequest(BaseModel):
    report_content: str
    job_title: str = ""
    company_name: str = ""

pdf_meta = {}

class HeaderBanner(Flowable):
    def __init__(self, width, height=60):
        Flowable.__init__(self)
        self.width = width
        self.height = height
    
    def draw(self):
        self.canv.setFillColor(HexColor('#1E3A5F'))
        self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)
        self.canv.setFillColor(white)
        self.canv.setFont('Helvetica-Bold', 18)
        self.canv.drawString(20, self.height - 25, "GetHiredAlly")
        self.canv.setFillColor(HexColor('#93C5FD'))
        self.canv.setFont('Helvetica', 12)
        self.canv.drawString(20, self.height - 45, "Job Analysis Report")

class InfoBox(Flowable):
    def __init__(self, width, job_title, company_name, height=60):
        Flowable.__init__(self)
        self.width = width
        self.height = height
        self.job_title = job_title
        self.company_name = company_name
    
    def draw(self):
        self.canv.setFillColor(HexColor('#F3F4F6'))
        self.canv.setStrokeColor(HexColor('#E5E7EB'))
        self.canv.setLineWidth(1)
        self.canv.roundRect(0, 0, self.width, self.height, 4, fill=1, stroke=1)
        
        y_pos = self.height - 18
        self.canv.setFillColor(HexColor('#000000'))
        self.canv.setFont('Helvetica-Bold', 10)
        self.canv.drawString(12, y_pos, f"Job Title: {self.job_title}")
        
        if self.company_name:
            y_pos -= 16
            self.canv.setFont('Helvetica', 10)
            self.canv.drawString(12, y_pos, f"Company: {self.company_name}")
        
        y_pos -= 16
        self.canv.setFillColor(HexColor('#6B7280'))
        self.canv.setFont('Helvetica', 10)
        self.canv.drawString(12, y_pos, f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

class SectionDivider(Flowable):
    def __init__(self, width):
        Flowable.__init__(self)
        self.width = width
        self.height = 1
    
    def draw(self):
        self.canv.setStrokeColor(HexColor('#E5E7EB'))
        self.canv.setLineWidth(0.5)
        self.canv.line(0, 0, self.width, 0)

class CalloutBox(Flowable):
    def __init__(self, width, text, box_type='info'):
        Flowable.__init__(self)
        self.width = width
        self.text = text
        self.box_type = box_type
        self.height = max(40, len(text) // 80 * 14 + 30)
    
    def draw(self):
        colors_map = {
            'red': {'bg': '#FEE2E2', 'border': '#EF4444'},
            'blue': {'bg': '#DBEAFE', 'border': '#3B82F6'},
            'green': {'bg': '#D1FAE5', 'border': '#10B981'}
        }
        c = colors_map.get(self.box_type, colors_map['blue'])
        
        self.canv.setFillColor(HexColor(c['bg']))
        self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)
        self.canv.setFillColor(HexColor(c['border']))
        self.canv.rect(0, 0, 4, self.height, fill=1, stroke=0)
        
        self.canv.setFillColor(HexColor('#000000'))
        self.canv.setFont('Helvetica', 10)
        
        words = self.text.split()
        lines = []
        current_line = ""
        for word in words:
            test_line = current_line + " " + word if current_line else word
            if self.canv.stringWidth(test_line, 'Helvetica', 10) < self.width - 24:
                current_line = test_line
            else:
                lines.append(current_line)
                current_line = word
        if current_line:
            lines.append(current_line)
        
        y_pos = self.height - 14
        for line in lines[:3]:
            self.canv.drawString(12, y_pos, line)
            y_pos -= 14

def add_first_page_footer(canvas, doc):
    canvas.saveState()
    canvas.setStrokeColor(HexColor('#E5E7EB'))
    canvas.setLineWidth(0.5)
    canvas.line(50, 45, A4[0] - 50, 45)
    canvas.setFont('Helvetica', 9)
    canvas.setFillColor(HexColor('#6B7280'))
    canvas.drawString(50, 30, "GetHiredAlly.com")
    url = "https://GetHiredAlly.com"
    url_width = canvas.stringWidth(url, 'Helvetica', 9)
    center_x = A4[0] / 2 - url_width / 2
    canvas.setFillColor(HexColor('#1E3A5F'))
    canvas.drawString(center_x, 30, url)
    canvas.linkURL(url, (center_x, 28, center_x + url_width, 40), relative=0)
    canvas.setFillColor(HexColor('#6B7280'))
    total_pages = pdf_meta.get('total_pages', doc.page)
    canvas.drawRightString(A4[0] - 50, 30, f"Page {doc.page} of {total_pages}")
    canvas.restoreState()

def add_later_page_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(HexColor('#1E3A5F'))
    canvas.rect(0, A4[1] - 35, A4[0], 35, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont('Helvetica', 10)
    canvas.drawString(50, A4[1] - 23, "GetHiredAlly - Job Analysis Report")
    canvas.setStrokeColor(HexColor('#E5E7EB'))
    canvas.setLineWidth(0.5)
    canvas.line(50, 45, A4[0] - 50, 45)
    canvas.setFont('Helvetica', 9)
    canvas.setFillColor(HexColor('#6B7280'))
    canvas.drawString(50, 30, "GetHiredAlly.com")
    url = "https://GetHiredAlly.com"
    url_width = canvas.stringWidth(url, 'Helvetica', 9)
    center_x = A4[0] / 2 - url_width / 2
    canvas.setFillColor(HexColor('#1E3A5F'))
    canvas.drawString(center_x, 30, url)
    canvas.linkURL(url, (center_x, 28, center_x + url_width, 40), relative=0)
    canvas.setFillColor(HexColor('#6B7280'))
    total_pages = pdf_meta.get('total_pages', doc.page)
    canvas.drawRightString(A4[0] - 50, 30, f"Page {doc.page} of {total_pages}")
    canvas.restoreState()

def parse_inline_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'`(.+?)`', r'<font face="Courier" size="9">\1</font>', text)
    return text

@router.post("/download/pdf")
async def download_pdf(request: DownloadRequest):
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=60
        )
        
        content_width = A4[0] - 100
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=26,
            spaceAfter=12,
            spaceBefore=16,
            textColor=HexColor('#1E3A5F'),
            fontName='Helvetica-Bold',
            alignment=TA_CENTER
        )
        
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=16,
            spaceAfter=8,
            textColor=HexColor('#1E3A5F'),
            fontName='Helvetica-Bold'
        )
        
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
            textColor=HexColor('#374151'),
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=15.4,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            leading=15.4,
            leftIndent=15,
            bulletIndent=0,
            fontName='Helvetica'
        )
        
        story = []
        
        story.append(HeaderBanner(content_width))
        story.append(Spacer(1, 16))
        
        title_text = request.job_title if request.job_title else "Job Analysis Report"
        info_height = 60 if request.company_name else 44
        story.append(InfoBox(content_width, title_text, request.company_name, info_height))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph(title_text, title_style))
        story.append(Spacer(1, 8))
        
        lines = request.report_content.split('\n')
        skip_separators = True
        
        for line in lines:
            line = line.strip()
            
            if line in ['---', '***', '___']:
                continue
            
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            if line.startswith('###'):
                clean_line = line.lstrip('#').strip()
                story.append(Paragraph(f"▪ {clean_line}", h3_style))
            elif line.startswith('##'):
                clean_line = line.lstrip('#').strip()
                story.append(SectionDivider(content_width))
                story.append(Spacer(1, 8))
                story.append(Paragraph(f"▪ {clean_line}", h2_style))
            elif line.startswith('#'):
                clean_line = line.lstrip('#').strip()
                story.append(SectionDivider(content_width))
                story.append(Spacer(1, 8))
                story.append(Paragraph(clean_line, h2_style))
            elif line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
                clean_line = line[2:]
                clean_line = parse_inline_markdown(clean_line)
                story.append(Paragraph(f'<font color="#1E3A5F">•</font> {clean_line}', bullet_style))
            elif '🚩' in line.lower() or 'red flag' in line.lower() or 'warning' in line.lower():
                clean_line = parse_inline_markdown(line)
                story.append(Spacer(1, 8))
                story.append(CalloutBox(content_width, line.replace('**', '').replace('*', ''), 'red'))
                story.append(Spacer(1, 8))
            elif 'insight' in line.lower() or 'tip' in line.lower() or 'note:' in line.lower():
                clean_line = parse_inline_markdown(line)
                story.append(Spacer(1, 8))
                story.append(CalloutBox(content_width, line.replace('**', '').replace('*', ''), 'blue'))
                story.append(Spacer(1, 8))
            elif 'success' in line.lower() or 'strength' in line.lower() or 'positive' in line.lower():
                clean_line = parse_inline_markdown(line)
                story.append(Spacer(1, 8))
                story.append(CalloutBox(content_width, line.replace('**', '').replace('*', ''), 'green'))
                story.append(Spacer(1, 8))
            elif line.startswith('**') and line.endswith('**'):
                clean_line = line.strip('*').strip()
                story.append(Paragraph(f"<b>{clean_line}</b>", body_style))
            else:
                clean_line = parse_inline_markdown(line)
                if clean_line:
                    story.append(Paragraph(clean_line, body_style))
        
        temp_buffer = io.BytesIO()
        temp_doc = SimpleDocTemplate(temp_buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=60)
        temp_doc.build(story.copy())
        pdf_meta['total_pages'] = temp_doc.page
        
        doc.build(story, onFirstPage=add_first_page_footer, onLaterPages=add_later_page_header_footer)
        buffer.seek(0)
        
        safe_company = request.company_name.replace(' ', '_').replace('/', '_') if request.company_name else 'Analysis'
        filename = f"XRay_Analysis_{safe_company}.pdf"
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

def set_run_font(run, font_name='Arial', size=11, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1.15):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    spacing.set(qn('w:line'), str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_bottom_border(paragraph, color_hex='1E3A5F', size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

@router.post("/download/docx")
async def download_docx(request: DownloadRequest):
    try:
        document = Document()
        
        navy_color = RGBColor(0x1E, 0x3A, 0x5F)
        dark_gray = RGBColor(0x37, 0x41, 0x51)
        white_color = RGBColor(0xFF, 0xFF, 0xFF)
        
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run1 = footer_para.add_run("GetHiredAlly.com  |  ")
            run1.font.name = 'Arial'
            run1.font.size = Pt(9)
            run1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            
            from docx.oxml.shared import OxmlElement as SharedOxmlElement
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), footer_para.part.relate_to(
                'https://GetHiredAlly.com',
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            ))
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '1E3A5F')
            rPr.append(color)
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            new_run.append(rPr)
            text_elem = OxmlElement('w:t')
            text_elem.text = 'https://GetHiredAlly.com'
            new_run.append(text_elem)
            hyperlink.append(new_run)
            footer_para._p.append(hyperlink)
            
            run2 = footer_para.add_run("  |  Page ")
            run2.font.name = 'Arial'
            run2.font.size = Pt(9)
            run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run_page = footer_para.add_run()
            run_page._r.append(fldChar1)
            run_page._r.append(instrText)
            run_page._r.append(fldChar2)
            run_page.font.name = 'Arial'
            run_page.font.size = Pt(9)
            run_page.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        
        header_table = document.add_table(rows=1, cols=1)
        header_table.autofit = False
        header_table.allow_autofit = False
        header_cell = header_table.cell(0, 0)
        header_cell.width = Inches(7)
        set_cell_shading(header_cell, '1E3A5F')
        
        header_para = header_cell.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run("GetHiredAlly - Job Analysis Report")
        set_run_font(header_run, size=14, bold=True, color=white_color)
        
        from docx.shared import Twips
        tc = header_cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'bottom', 'left', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '200')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        document.add_paragraph()
        
        info_table = document.add_table(rows=1, cols=1)
        info_table.autofit = False
        info_cell = info_table.cell(0, 0)
        info_cell.width = Inches(7)
        set_cell_shading(info_cell, 'F3F4F6')
        
        tc_info = info_cell._tc
        tcPr_info = tc_info.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'E5E7EB')
            tcBorders.append(border)
        tcPr_info.append(tcBorders)
        
        tcMar_info = OxmlElement('w:tcMar')
        for margin_name in ['top', 'bottom', 'left', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '150')
            margin.set(qn('w:type'), 'dxa')
            tcMar_info.append(margin)
        tcPr_info.append(tcMar_info)
        
        info_para = info_cell.paragraphs[0]
        title_text = request.job_title if request.job_title else "Job Analysis Report"
        job_run = info_para.add_run(f"Job Title: {title_text}")
        set_run_font(job_run, size=10, bold=True)
        
        if request.company_name:
            info_para.add_run("\n")
            company_run = info_para.add_run(f"Company: {request.company_name}")
            set_run_font(company_run, size=10)
        
        info_para.add_run("\n")
        date_run = info_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        set_run_font(date_run, size=10, color=dark_gray)
        
        document.add_paragraph()
        
        title_para = document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title_text)
        set_run_font(title_run, size=26, bold=True, color=navy_color)
        set_paragraph_spacing(title_para, space_after=12)
        
        lines = request.report_content.split('\n')
        skip_next_empty = False
        
        for line in lines:
            line = line.strip()
            
            if line == '---' or line == '***' or line == '___':
                skip_next_empty = True
                continue
            
            if not line:
                if not skip_next_empty:
                    spacer = document.add_paragraph()
                    set_paragraph_spacing(spacer, space_before=2, space_after=2)
                skip_next_empty = False
                continue
            
            skip_next_empty = False
            
            if line.startswith('###'):
                clean_line = line.lstrip('#').strip()
                p = document.add_paragraph()
                run = p.add_run(clean_line)
                set_run_font(run, size=12, bold=True, color=dark_gray)
                set_paragraph_spacing(p, space_before=10, space_after=6)
            elif line.startswith('##'):
                clean_line = line.lstrip('#').strip()
                p = document.add_paragraph()
                add_bottom_border(p, '1E3A5F', 4)
                run = p.add_run(clean_line)
                set_run_font(run, size=14, bold=True, color=navy_color)
                set_paragraph_spacing(p, space_before=16, space_after=8)
            elif line.startswith('#'):
                clean_line = line.lstrip('#').strip()
                p = document.add_paragraph()
                add_bottom_border(p, '1E3A5F', 6)
                run = p.add_run(clean_line)
                set_run_font(run, size=16, bold=True, color=navy_color)
                set_paragraph_spacing(p, space_before=16, space_after=10)
            elif line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
                clean_line = line[2:].replace('**', '').replace('*', '')
                p = document.add_paragraph(style='List Bullet')
                run = p.add_run(clean_line)
                set_run_font(run, size=11)
                set_paragraph_spacing(p, space_before=1, space_after=1)
            elif line.startswith('**') and line.endswith('**'):
                clean_line = line.strip('*').strip()
                p = document.add_paragraph()
                run = p.add_run(clean_line)
                set_run_font(run, size=11, bold=True)
                set_paragraph_spacing(p, space_before=4, space_after=4)
            else:
                clean_line = line.replace('**', '').replace('*', '')
                if clean_line:
                    p = document.add_paragraph()
                    run = p.add_run(clean_line)
                    set_run_font(run, size=11)
                    set_paragraph_spacing(p, space_before=2, space_after=4)
        
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        safe_company = request.company_name.replace(' ', '_').replace('/', '_') if request.company_name else 'Analysis'
        filename = f"XRay_Analysis_{safe_company}.docx"
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word document generation failed: {str(e)}")
