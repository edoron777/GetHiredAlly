from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
import logging

logger = logging.getLogger(__name__)


def generate_pdf(content: str) -> bytes:
    """Generate PDF from text content."""
    logger.info(f"[PDF_GEN] Starting PDF generation, content length: {len(content)} chars")
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=11)
    pdf.set_auto_page_break(auto=True, margin=15)

    lines = content.split('\n')
    logger.info(f"[PDF_GEN] Processing {len(lines)} lines")

    for i, line in enumerate(lines):
        try:
            if line.isupper() and len(line) < 50 and line.strip():
                pdf.set_font('Arial', 'B', 12)
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.cell(0, 8, safe_line, ln=True)
                pdf.set_font('Arial', size=11)
            else:
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, safe_line)
        except Exception as e:
            logger.warning(f"[PDF_GEN] Error on line {i}: {str(e)}, skipping line")
            continue

    logger.info("[PDF_GEN] Generating output bytes")
    
    try:
        output = pdf.output(dest='S')
        if isinstance(output, str):
            result = output.encode('latin-1')
        else:
            result = bytes(output)
        logger.info(f"[PDF_GEN] PDF generated successfully, size: {len(result)} bytes")
        return result
    except Exception as e:
        logger.error(f"[PDF_GEN] Output generation failed: {type(e).__name__}: {str(e)}")
        raise


def generate_docx(content: str) -> bytes:
    """Generate DOCX from text content."""
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')

    for line in lines:
        if not line.strip():
            doc.add_paragraph()
        elif line.isupper() and len(line) < 50:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(14)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
